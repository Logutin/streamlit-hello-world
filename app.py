import streamlit as st
import docx
import pandas as pd
import io

# --- Function to Extract Last Table (Updated) ---
def extract_last_table_as_df(uploaded_file):
    """
    Reads a DOCX file, extracts the last table, cleans headers for uniqueness,
    and returns it as a Pandas DataFrame.
    """
    try:
        document = docx.Document(uploaded_file)

        if not document.tables:
            return None # No tables found

        last_table = document.tables[-1]
        data = []
        for row in last_table.rows:
            # Ensure cell text is stripped of leading/trailing whitespace
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)

        if not data:
            return None # Table was empty

        # --- Header Cleaning Logic ---
        raw_headers = data[0]
        df_data = data[1:]
        processed_headers = []
        counts = {} # To track counts of each header name seen so far

        for i, header in enumerate(raw_headers):
            # If header is empty or None, create a default name
            if not header:
                col_name = f"Unnamed_{i+1}" # Use 1-based index for user display
            else:
                col_name = header # Use the original name initially

            # Check for duplicates and append count if needed
            if col_name in counts:
                counts[col_name] += 1
                processed_headers.append(f"{col_name}_{counts[col_name]}")
            else:
                counts[col_name] = 1
                processed_headers.append(col_name)
        # --- End Header Cleaning ---

        # Create DataFrame with cleaned headers
        df = pd.DataFrame(df_data, columns=processed_headers)
        return df

    except Exception as e:
        # More specific error logging can be helpful
        st.error(f"Error processing DOCX table: {e}")
        # You could log the full traceback here if needed for debugging
        # import traceback
        # st.error(traceback.format_exc())
        return None
# --- End of Function ---


# --- Streamlit App (Main part remains largely the same) ---
st.title("DOCX Uploader, Reader, and Table Extractor")

st.write("Upload a .docx file. We'll show details and extract the last table.")

uploaded_file = st.file_uploader("Choose a DOCX file", type=["docx"])

if uploaded_file is not None:
    st.success(f"Successfully uploaded: {uploaded_file.name}")

    # --- Basic File Details ---
    st.write("---")
    st.write("### Basic File Details:")
    st.write(f"**Name:** {uploaded_file.name}")
    st.write(f"**Type:** {uploaded_file.type}")
    st.write(f"**Size:** {uploaded_file.size} bytes")
    st.write("---")

    # --- Document Content Analysis ---
    st.write("### Document Content Analysis:")
    try:
        # Reset buffer position before reading again is crucial
        uploaded_file.seek(0)
        document = docx.Document(uploaded_file)
        num_paragraphs = len(document.paragraphs)
        num_tables = len(document.tables)
        st.write(f"**Number of paragraphs:** {num_paragraphs}")
        st.write(f"**Number of tables:** {num_tables}")

    except Exception as e:
        st.error(f"Error reading basic DOCX properties: {e}")

    st.write("---")

    # --- Extract and Display Last Table ---
    st.write("### Last Table Extraction:")

    # Reset buffer position again before passing to the table extraction function
    uploaded_file.seek(0)
    df_last_table = extract_last_table_as_df(uploaded_file)

    if df_last_table is not None:
        if not df_last_table.empty:
            st.info("Displaying the last table found. Headers were automatically cleaned for uniqueness if necessary.") # Added info
            st.dataframe(df_last_table)
        else:
            st.warning("The last table found in the document appears to be empty (no data rows).") # Clarified message
    else:
        st.warning("Could not find any tables in the document or failed to read the last one.")

    st.write("---")
    st.info("Note: File processed in memory, not saved permanently.")

else:
    st.markdown("👆 Upload a file using the button above.")