# utils.py
import pandas as pd
import docx
import io
import streamlit as st # Import Streamlit for error display if needed

def clean_and_uniquify_headers(raw_headers):
    """Takes a list of raw header strings and makes them unique and non-empty."""
    processed_headers = []
    counts = {}
    for i, header in enumerate(raw_headers):
        header = str(header).strip() # Ensure string and strip whitespace
        if not header:
            col_name = f"Unnamed_{i+1}"
        else:
            col_name = header

        if col_name in counts:
            counts[col_name] += 1
            processed_headers.append(f"{col_name}_{counts[col_name]}")
        else:
            counts[col_name] = 1
            processed_headers.append(col_name)
    return processed_headers

def extract_last_table_as_df(file_bytes, use_first_row_as_header=True):
    """
    Reads DOCX file bytes, extracts the last table into a Pandas DataFrame.

    Args:
        file_bytes (bytes): The content of the uploaded DOCX file.
        use_first_row_as_header (bool): If True, use the first row as headers.
                                        Otherwise, generate default headers.

    Returns:
        pandas.DataFrame or None: The extracted table or None on failure/no table.
    """
    try:
        # Use io.BytesIO to treat the bytes object like a file
        file_buffer = io.BytesIO(file_bytes)
        document = docx.Document(file_buffer)

        if not document.tables:
            return None # No tables found

        last_table = document.tables[-1]
        data = []
        for row in last_table.rows:
            row_data = [cell.text for cell in row.cells]
            data.append(row_data)

        if not data:
            return None # Table was empty

        df = None
        if use_first_row_as_header:
            if len(data) > 0: # Need at least one row for headers
                raw_headers = data[0]
                df_data = data[1:]
                processed_headers = clean_and_uniquify_headers(raw_headers)
                df = pd.DataFrame(df_data, columns=processed_headers)
            else:
                # Table only had a header row? Return empty df with cleaned headers
                processed_headers = clean_and_uniquify_headers(data[0])
                df = pd.DataFrame(columns=processed_headers)
        else:
            # Don't use the first row as header, generate default headers
            num_columns = len(data[0]) if data else 0
            default_headers = [f"Column_{i+1}" for i in range(num_columns)]
            # All extracted rows are data
            df = pd.DataFrame(data, columns=default_headers)

        return df

    except Exception as e:
        # Log the error appropriately (e.g., using st.error if called from Streamlit context,
        # or standard logging if used elsewhere)
        # For now, just print and return None
        print(f"Error processing DOCX table: {e}")
        # Optionally display error in Streamlit if possible, though utils shouldn't depend on st UI directly
        # st.error(f"Error processing DOCX table: {e}")
        return None